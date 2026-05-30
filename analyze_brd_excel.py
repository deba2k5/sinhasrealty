import openpyxl
import pandas as pd
import json
from docx import Document
import os

# Analyze Excel file
excel_file = r"c:\Users\Debangshu05\Downloads\sinharealty\SINHAS_Own_Property_Purchase_Mortgage_Register-2 (1).xlsx"
docx_file = r"c:\Users\Debangshu05\Downloads\sinharealty\SINHAS_BRS_Own_Property_Purchase_Mortgage_Register-2 (1).docx"

print("=" * 80)
print("EXCEL FILE ANALYSIS")
print("=" * 80)

try:
    xls = pd.ExcelFile(excel_file, engine='openpyxl')
    print(f"\nSheet names: {xls.sheet_names}\n")
    
    for sheet in xls.sheet_names:
        print(f"\n--- Sheet: {sheet} ---")
        df = pd.read_excel(excel_file, sheet_name=sheet)
        print(f"Shape: {df.shape}")
        print(f"Columns: {list(df.columns)}")
        print(f"First few rows:")
        print(df.head(2).to_string())
        print(f"Data types:\n{df.dtypes}")
        
except Exception as e:
    print(f"Error reading Excel: {e}")

print("\n" + "=" * 80)
print("BRD DOCUMENT ANALYSIS")
print("=" * 80)

try:
    doc = Document(docx_file)
    print(f"\nBRD Document Content:\n")
    for para in doc.paragraphs:
        if para.text.strip():
            print(para.text)
    
    print("\n\nTables in BRD:")
    for i, table in enumerate(doc.tables):
        print(f"\n--- Table {i+1} ---")
        for row in table.rows:
            print([cell.text for cell in row.cells])
            
except Exception as e:
    print(f"Error reading DOCX: {e}")
